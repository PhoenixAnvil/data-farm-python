"""One-off script to build CODE_REVIEW.docx from CODE_REVIEW.md. Run: python build_review_docx.py"""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def main():
    md_path = Path(__file__).resolve().parent / "CODE_REVIEW.md"
    docx_path = Path(__file__).resolve().parent / "CODE_REVIEW.docx"
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        stripped = line.rstrip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("- [ ] "):
            p = doc.add_paragraph(stripped[6:], style="List Bullet")
            p.paragraph_format.left_indent = Pt(24)
        elif stripped.startswith("- [x] "):
            p = doc.add_paragraph(stripped[6:] + " (done)", style="List Bullet")
            p.paragraph_format.left_indent = Pt(24)
        elif stripped.startswith("- **") and ":**" in stripped:
            doc.add_paragraph(stripped, style="List Bullet")
        elif stripped in ("---",):
            doc.add_paragraph("—" * 20)
        else:
            doc.add_paragraph(stripped)

    doc.save(docx_path)
    print(f"Saved {docx_path}")

if __name__ == "__main__":
    main()
